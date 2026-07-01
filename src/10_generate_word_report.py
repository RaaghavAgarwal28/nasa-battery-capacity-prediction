"""
=========================================================
NASA Battery Capacity Prediction
Generate Final Report Word Document
=========================================================
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# =========================================================
# PATHS
# =========================================================
PROJECT_ROOT  = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_FOLDER = os.path.join(PROJECT_ROOT, "outputs")
GRAPH_FOLDER  = os.path.join(OUTPUT_FOLDER, "report_graphs")
REPORT_PATH   = os.path.join(PROJECT_ROOT, "reports", "NASA_Battery_Report.docx")
os.makedirs(os.path.join(PROJECT_ROOT, "reports"), exist_ok=True)

# =========================================================
# HELPERS
# =========================================================

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   'single')
        tag.set(qn('w:sz'),    '6')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), 'CCCCCC')
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color='1F3864'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    sizes = {1: 22, 2: 16, 3: 13}
    run.font.size = Pt(sizes.get(level, 13))
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_body(doc, text, bold=False, italic=False, size=11, color='222222'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_image_centered(doc, img_path, width_inches=6.0):
    if not os.path.exists(img_path):
        doc.add_paragraph(f"[Image not found: {img_path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))

def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string('666666')

def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_metrics_table(doc, rows, header_bg='1F3864', row_bg_alt='EBF0F8'):
    """rows = list of lists. First row = headers."""
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_text))
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(11)
                set_cell_bg(cell, header_bg)
            else:
                run.font.size = Pt(10.5)
                if r_idx % 2 == 0:
                    set_cell_bg(cell, row_bg_alt)
                else:
                    set_cell_bg(cell, 'FFFFFF')
            set_cell_border(cell)
    return table

def add_page_break(doc):
    doc.add_page_break()

# =========================================================
# CREATE DOCUMENT
# =========================================================
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# =========================================================
# COVER PAGE
# =========================================================
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("NASA Lithium-Ion Battery Ageing")
title_run.bold = True
title_run.font.size = Pt(26)
title_run.font.color.rgb = RGBColor.from_string('1F3864')

title_p2 = doc.add_paragraph()
title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run2 = title_p2.add_run("Capacity Prediction — Model Report")
title_run2.bold = True
title_run2.font.size = Pt(22)
title_run2.font.color.rgb = RGBColor.from_string('2E74B5')

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run("A Regression Study using Linear Regression, Random Forest & LSTM")
sub_run.italic = True
sub_run.font.size = Pt(13)
sub_run.font.color.rgb = RGBColor.from_string('444444')

doc.add_paragraph()
doc.add_paragraph()

add_image_centered(doc, os.path.join(GRAPH_FOLDER, "fig7_dashboard.png"), width_inches=6.3)
add_caption(doc, "Fig. Grand Summary Dashboard — All Models")

doc.add_paragraph()
auth_p = doc.add_paragraph()
auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth_run = auth_p.add_run("Raaghav Agarwal  ·  NASA Battery Aging Dataset")
auth_run.font.size = Pt(11)
auth_run.font.color.rgb = RGBColor.from_string('666666')

add_page_break(doc)

# =========================================================
# SLIDE 1 — LINEAR REGRESSION
# =========================================================
add_heading(doc, "1.  Linear Regression", level=1)
add_body(doc, "Baseline Model — Phase 5", italic=True, color='555555', size=11)
add_divider(doc)
doc.add_paragraph()

add_heading(doc, "What is Linear Regression?", level=2)
for bullet in [
    "Models capacity as a linear combination of 17 engineered features",
    "Simple, interpretable, and computationally inexpensive",
    "Serves as a robust baseline for comparison with more complex models",
]:
    p = doc.add_paragraph(bullet, style='List Bullet')
    p.runs[0].font.size = Pt(11)

doc.add_paragraph()
add_heading(doc, "Training Setup", level=2)
add_metrics_table(doc, [
    ["Parameter", "Value"],
    ["Features",        "17 engineered cycle statistics"],
    ["Train / Test",    "80% / 20%"],
    ["Scaling",         "StandardScaler (zero mean, unit variance)"],
    ["Library",         "scikit-learn LinearRegression"],
])

doc.add_paragraph()
add_heading(doc, "Results", level=2)
add_metrics_table(doc, [
    ["Metric", "Value"],
    ["MAE",  "0.1029 Ah"],
    ["RMSE", "0.1523 Ah"],
    ["R²",   "0.8874"],
], header_bg='C00000')

doc.add_paragraph()
add_heading(doc, "Key Observations", level=2)
for bullet in [
    "Achieves R² ≈ 0.89 — a solid baseline but underfits non-linear degradation patterns",
    "Residuals show mild heteroscedasticity — prediction errors grow at lower capacities",
    "Outlier filtering of Re and Rct impedance sensors was critical for numerical stability",
]:
    p = doc.add_paragraph(bullet, style='List Bullet')
    p.runs[0].font.size = Pt(11)

doc.add_paragraph()
add_heading(doc, "Plots", level=2)

add_image_centered(doc, os.path.join(GRAPH_FOLDER, "fig1_actual_vs_predicted.png"), width_inches=6.3)
add_caption(doc, "Fig 1. Actual vs Predicted Capacity — Linear Regression (left panel)")

doc.add_paragraph()
add_image_centered(doc, os.path.join(GRAPH_FOLDER, "fig2_prediction_error_scatter.png"), width_inches=6.3)
add_caption(doc, "Fig 2. Prediction Error Scatter — Linear Regression (left panel)")

doc.add_paragraph()
add_image_centered(doc, os.path.join(GRAPH_FOLDER, "fig3_residuals.png"), width_inches=6.3)
add_caption(doc, "Fig 3. Residual Analysis — Linear Regression (left panel)")

add_page_break(doc)

# =========================================================
# SLIDE 2 — RANDOM FOREST
# =========================================================
add_heading(doc, "2.  Random Forest", level=1)
add_body(doc, "Ensemble Tree Model — Phase 6  🏆 Best Performer", italic=True, color='555555', size=11)
add_divider(doc)
doc.add_paragraph()

add_heading(doc, "What is Random Forest?", level=2)
for bullet in [
    "Ensemble of 100 decision trees trained on random subsets of data and features",
    "Captures non-linear relationships and feature interactions automatically",
    "Robust to overfitting via bagging (Bootstrap Aggregation)",
]:
    p = doc.add_paragraph(bullet, style='List Bullet')
    p.runs[0].font.size = Pt(11)

doc.add_paragraph()
add_heading(doc, "Training Setup", level=2)
add_metrics_table(doc, [
    ["Parameter", "Value"],
    ["Features",      "17 engineered cycle statistics (flat, non-sequential)"],
    ["Train / Test",  "80% / 20%"],
    ["Estimators",    "100 trees"],
    ["Max Features",  "sqrt(n_features)"],
    ["Library",       "scikit-learn RandomForestRegressor"],
])

doc.add_paragraph()
add_heading(doc, "Results", level=2)
add_metrics_table(doc, [
    ["Metric", "Value"],
    ["MAE",  "0.0142 Ah  ✅ Best"],
    ["RMSE", "0.0585 Ah  ✅ Best"],
    ["R²",   "0.9834     ✅ Best"],
], header_bg='375623')

doc.add_paragraph()
add_heading(doc, "Key Observations", level=2)
for bullet in [
    "Best performing model with R² of 0.983 — captures the degradation curve with high precision",
    "Predictions tightly cluster around the perfect-fit diagonal line",
    "Top predictive features: Cycle_Number, Discharge_Duration, Voltage_Mean",
    "Recommended model for production deployment",
]:
    p = doc.add_paragraph(bullet, style='List Bullet')
    p.runs[0].font.size = Pt(11)

doc.add_paragraph()
add_heading(doc, "Plots", level=2)

add_image_centered(doc, os.path.join(GRAPH_FOLDER, "fig1_actual_vs_predicted.png"), width_inches=6.3)
add_caption(doc, "Fig 1. Actual vs Predicted Capacity — Random Forest (middle panel)")

doc.add_paragraph()
add_image_centered(doc, os.path.join(GRAPH_FOLDER, "fig2_prediction_error_scatter.png"), width_inches=6.3)
add_caption(doc, "Fig 2. Prediction Error Scatter — Random Forest (middle panel)")

doc.add_paragraph()
add_image_centered(doc, os.path.join(GRAPH_FOLDER, "fig6_feature_importance.png"), width_inches=6.3)
add_caption(doc, "Fig 6. Feature Importances — Random Forest")

add_page_break(doc)

# =========================================================
# SLIDE 3 — LSTM
# =========================================================
add_heading(doc, "3.  LSTM Neural Network", level=1)
add_body(doc, "Deep Learning Sequential Model — Phase 7", italic=True, color='555555', size=11)
add_divider(doc)
doc.add_paragraph()

add_heading(doc, "What is LSTM?", level=2)
for bullet in [
    "Long Short-Term Memory — a type of Recurrent Neural Network (RNN)",
    "Designed to model temporal sequences — captures how capacity degrades over cycles",
    "Uses a sliding window of 5 past cycles to predict the current cycle's capacity",
]:
    p = doc.add_paragraph(bullet, style='List Bullet')
    p.runs[0].font.size = Pt(11)

doc.add_paragraph()
add_heading(doc, "Network Architecture", level=2)
add_metrics_table(doc, [
    ["Layer",          "Configuration"],
    ["Input",          "Shape: (5 cycles × 17 features)"],
    ["LSTM Layer 1",   "64 units, return_sequences=True"],
    ["Dropout 1",      "Rate = 0.2"],
    ["LSTM Layer 2",   "32 units, return_sequences=False"],
    ["Dropout 2",      "Rate = 0.2"],
    ["Dense (hidden)", "16 units, ReLU activation"],
    ["Dense (output)", "1 unit (Capacity)"],
    ["Optimizer",      "Adam"],
    ["Loss",           "Mean Squared Error (MSE)"],
    ["Early Stopping", "Patience = 10 epochs"],
])

doc.add_paragraph()
add_heading(doc, "Training Setup", level=2)
add_metrics_table(doc, [
    ["Parameter",       "Value"],
    ["Sequence length",  "5 cycles (sliding window)"],
    ["Batch size",       "16"],
    ["Max epochs",       "60"],
    ["Validation split", "15% of training data"],
    ["Library",          "TensorFlow / Keras"],
])

doc.add_paragraph()
add_heading(doc, "Results", level=2)
add_metrics_table(doc, [
    ["Metric", "Value"],
    ["MAE",  "0.0784 Ah"],
    ["RMSE", "0.1169 Ah"],
    ["R²",   "0.9337"],
], header_bg='1F3864')

doc.add_paragraph()
add_heading(doc, "Key Observations", level=2)
for bullet in [
    "Explicitly models temporal degradation trajectory — understands trends across cycles",
    "Loss curve shows stable convergence; early stopping prevents overfitting",
    "Best suited for multi-step ahead forecasting tasks",
    "Slightly lower R² than Random Forest due to smaller effective training set after windowing",
]:
    p = doc.add_paragraph(bullet, style='List Bullet')
    p.runs[0].font.size = Pt(11)

doc.add_paragraph()
add_heading(doc, "Plots", level=2)

add_image_centered(doc, os.path.join(GRAPH_FOLDER, "fig1_actual_vs_predicted.png"), width_inches=6.3)
add_caption(doc, "Fig 1. Actual vs Predicted Capacity — LSTM (right panel)")

doc.add_paragraph()
add_image_centered(doc, os.path.join(GRAPH_FOLDER, "fig2_prediction_error_scatter.png"), width_inches=6.3)
add_caption(doc, "Fig 2. Prediction Error Scatter — LSTM (right panel)")

doc.add_paragraph()
add_image_centered(doc, os.path.join(GRAPH_FOLDER, "fig5_lstm_loss_curve.png"), width_inches=5.5)
add_caption(doc, "Fig 5. LSTM Training & Validation Loss Curve")

add_page_break(doc)

# =========================================================
# SLIDE 4 — COMPARISON
# =========================================================
add_heading(doc, "4.  Model Comparison & Conclusion", level=1)
add_body(doc, "Linear Regression vs Random Forest vs LSTM", italic=True, color='555555', size=11)
add_divider(doc)
doc.add_paragraph()

add_heading(doc, "Performance Summary", level=2)
add_metrics_table(doc, [
    ["Model",               "MAE ↓",    "RMSE ↓",   "R² ↑",    "Rank"],
    ["Linear Regression",   "0.1029",   "0.1523",   "0.8874",  "🥉 3rd"],
    ["LSTM",                "0.0784",   "0.1169",   "0.9337",  "🥈 2nd"],
    ["Random Forest",       "0.0142",   "0.0585",   "0.9834",  "🥇 Best"],
], header_bg='1F3864')

doc.add_paragraph()
add_heading(doc, "Model Characteristics", level=2)
add_metrics_table(doc, [
    ["Aspect",                  "Linear Regression",  "Random Forest",  "LSTM"],
    ["Captures non-linearity",  "No",                 "Yes",            "Yes"],
    ["Uses temporal sequences", "No",                 "No",             "Yes"],
    ["Interpretable",           "Yes",                "Partial",        "No"],
    ["Training speed",          "Very fast",          "Moderate",       "Slow"],
    ["Best use case",           "Baseline",           "Production",     "Forecasting"],
])

doc.add_paragraph()
add_heading(doc, "Key Takeaways", level=2)
for bullet in [
    "Random Forest is the recommended model — highest R² (0.983), lowest MAE (0.014 Ah)",
    "Linear Regression provides an interpretable baseline with R² = 0.887",
    "LSTM captures temporal degradation patterns — best suited for multi-step forecasting",
    "Outlier filtering of impedance sensors (Re, Rct) was critical for stable Linear Regression training",
    "Cycle_Number and Discharge_Duration are the most predictive features of capacity fade",
]:
    p = doc.add_paragraph(bullet, style='List Bullet')
    p.runs[0].font.size = Pt(11)

doc.add_paragraph()
add_heading(doc, "Comparison Plots", level=2)

add_image_centered(doc, os.path.join(GRAPH_FOLDER, "fig4_metric_comparison.png"), width_inches=6.3)
add_caption(doc, "Fig 4. MAE, RMSE and R² Comparison Across All Models")

doc.add_paragraph()
add_image_centered(doc, os.path.join(GRAPH_FOLDER, "fig3_residuals.png"), width_inches=6.3)
add_caption(doc, "Fig 3. Residual Analysis — All Three Models")

doc.add_paragraph()
add_image_centered(doc, os.path.join(GRAPH_FOLDER, "fig7_dashboard.png"), width_inches=6.3)
add_caption(doc, "Fig 7. Grand Summary Dashboard — All Models")

# =========================================================
# SAVE
# =========================================================
doc.save(REPORT_PATH)
print(f"\nWord report saved to: {REPORT_PATH}")
